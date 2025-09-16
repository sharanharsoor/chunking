"""
DOC/DOCX/ODT/RTF document chunking strategy with comprehensive format support.

This module provides specialized chunking for Microsoft Word and other document formats:
- DOC/DOCX (Microsoft Word)
- ODT (OpenDocument Text)
- RTF (Rich Text Format)
- Multiple backend support for maximum compatibility
"""

import logging
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
import tempfile

try:
    from docx import Document as DocxDocument
    from docx.opc.exceptions import PackageNotFoundError
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    import mammoth
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False

try:
    from odf.opendocument import load as odf_load
    from odf.text import P as ODFParagraph
    from odf import teletype
    HAS_ODF = True
except ImportError:
    HAS_ODF = False

try:
    from striprtf.striprtf import rtf_to_text
    HAS_STRIPRTF = True
except ImportError:
    HAS_STRIPRTF = False

try:
    import textract
    HAS_TEXTRACT = True
except ImportError:
    HAS_TEXTRACT = False

from chunking_strategy.core.base import BaseChunker, Chunk, ChunkingResult, ChunkMetadata, ModalityType
from chunking_strategy.core.registry import register_chunker, ComplexityLevel, SpeedLevel, MemoryUsage
from chunking_strategy.core.streaming import StreamableChunker
from chunking_strategy.core.adaptive import AdaptableChunker


@register_chunker(
    name="doc_chunker",
    category="document",
    complexity=ComplexityLevel.MEDIUM,
    speed=SpeedLevel.MEDIUM,
    memory=MemoryUsage.MEDIUM,
    supported_formats=[".docx", ".doc", ".odt", ".rtf"],
    dependencies=[],  # No required dependencies since any backend can work
    optional_dependencies=["python-docx", "mammoth", "odfpy", "striprtf", "textract"],
    description="Document chunker for Word, OpenDocument, and RTF formats with multiple backend support",
    use_cases=["document_processing", "office_documents", "text_extraction", "report_analysis"],
    best_for=["structured_documents", "formatted_text", "office_workflows", "content_migration"],
    limitations=["complex_formatting_may_be_lost", "requires_appropriate_backend"]
)
class DocChunker(StreamableChunker, AdaptableChunker):
    """
    Specialized chunker for document formats supporting:
    - DOC/DOCX (Microsoft Word documents)
    - ODT (OpenDocument Text)
    - RTF (Rich Text Format)
    - Multiple chunking strategies (paragraph, page, section, heading-based)
    - Content preservation with formatting metadata
    """

    def __init__(
        self,
        chunk_by: str = "paragraphs",  # "paragraphs", "sections", "headings", "pages", "content_size"
        paragraphs_per_chunk: int = 5,
        preserve_formatting: bool = True,
        extract_images: bool = False,  # For future enhancement
        extract_tables: bool = True,
        heading_levels: Optional[List[int]] = None,  # Which heading levels to split on
        min_chunk_size: int = 100,
        max_chunk_size: int = 10000,
        chunk_overlap: int = 50,
        backend: str = "auto",  # "auto", "python-docx", "mammoth", "textract", "odf", "striprtf"
        **kwargs
    ):
        """
        Initialize document chunker.

        Args:
            chunk_by: Chunking strategy to use
            paragraphs_per_chunk: Number of paragraphs per chunk
            preserve_formatting: Whether to preserve formatting information
            extract_images: Whether to extract embedded images (future feature)
            extract_tables: Whether to extract and format tables
            heading_levels: Heading levels to use for splitting (1-6)
            min_chunk_size: Minimum chunk size in characters
            max_chunk_size: Maximum chunk size in characters
            chunk_overlap: Overlap between chunks in characters
            backend: Backend to use for document processing
            **kwargs: Additional parameters
        """
        super().__init__(
            name="doc_chunker",
            category="document",
            supported_modalities=[ModalityType.TEXT, ModalityType.TABLE],
            **kwargs
        )

        self.chunk_by = chunk_by
        self.paragraphs_per_chunk = paragraphs_per_chunk
        self.preserve_formatting = preserve_formatting
        self.extract_images = extract_images
        self.extract_tables = extract_tables
        self.heading_levels = heading_levels or [1, 2, 3]
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.chunk_overlap = chunk_overlap
        self.backend = backend

        self.logger = logging.getLogger(__name__)

        # Setup backend
        self._setup_backend()

    def _setup_backend(self) -> None:
        """Setup the document processing backend."""
        if self.backend == "auto":
            # Choose best available backend
            if HAS_PYTHON_DOCX:
                self.backend = "python-docx"
            elif HAS_MAMMOTH:
                self.backend = "mammoth"
            elif HAS_TEXTRACT:
                self.backend = "textract"
            elif HAS_ODF:
                self.backend = "odf"
            elif HAS_STRIPRTF:
                self.backend = "striprtf"
            else:
                error_msg = (
                    "\n🚨 Missing Document Processing Dependencies 🚨\n\n"
                    "The doc_chunker requires at least one document processing library to be installed.\n"
                    "Please install one or more of the following packages:\n\n"
                    "📄 For Word Documents (.docx, .doc):\n"
                    "   pip install python-docx    # Recommended for .docx files\n"
                    "   pip install mammoth        # Good for .doc and .docx files\n\n"
                    "📄 For Multiple Document Types:\n"
                    "   pip install textract       # Supports many formats (PDF, DOCX, RTF, etc.)\n\n"
                    "📄 For OpenDocument Format (.odt):\n"
                    "   pip install odfpy          # For OpenDocument files\n\n"
                    "📄 For RTF Files (.rtf):\n"
                    "   pip install striprtf       # For Rich Text Format files\n\n"
                    "💡 Quick Install (Recommended):\n"
                    "   pip install python-docx mammoth textract odfpy striprtf\n\n"
                    "After installation, restart your application and try again."
                )
                raise ImportError(error_msg)

        self.logger.info(f"Using document backend: {self.backend}")

    @classmethod
    def check_dependencies(cls) -> Dict[str, bool]:
        """Check which document processing dependencies are available."""
        return {
            "python-docx": HAS_PYTHON_DOCX,
            "mammoth": HAS_MAMMOTH,
            "textract": HAS_TEXTRACT,
            "odfpy": HAS_ODF,
            "striprtf": HAS_STRIPRTF
        }

    @classmethod
    def get_dependency_status(cls) -> str:
        """Get a human-readable status of dependencies."""
        deps = cls.check_dependencies()
        available = [name for name, installed in deps.items() if installed]
        missing = [name for name, installed in deps.items() if not installed]

        status = "📋 Document Chunker Dependency Status:\n\n"

        if available:
            status += "✅ Available:\n"
            for dep in available:
                status += f"   • {dep}\n"
            status += "\n"

        if missing:
            status += "❌ Missing:\n"
            for dep in missing:
                status += f"   • {dep}\n"
            status += "\n"
            status += "💡 Install missing dependencies with:\n"
            status += f"   pip install {' '.join(missing)}\n"
        else:
            status += "🎉 All dependencies are available!\n"

        return status

    def chunk(
        self,
        content: Union[str, bytes, Path],
        source_info: Optional[Dict[str, Any]] = None,
        **kwargs
    ) -> ChunkingResult:
        """
        Chunk document content using specified strategy.

        Args:
            content: Document file path or content bytes
            source_info: Source information metadata
            **kwargs: Additional chunking parameters

        Returns:
            ChunkingResult with document chunks
        """
        start_time = time.time()

        # Handle different input types
        if isinstance(content, Path):
            doc_path = content
        elif isinstance(content, str):
            # Check if it's a file path or actual content
            if (len(content) < 300 and '\n' not in content and
                content.strip() and not content.isspace()):
                try:
                    if Path(content).exists() and Path(content).is_file():
                        doc_path = Path(content)
                    else:
                        # Content string - save to temporary file
                        doc_path = self._save_content_to_temp_file(content)
                except (OSError, ValueError):
                    # If path check fails (e.g., filename too long), treat as content
                    doc_path = self._save_content_to_temp_file(content)
            else:
                # Content string - save to temporary file
                doc_path = self._save_content_to_temp_file(content)
        elif isinstance(content, bytes):
            # Save bytes to temporary file
            doc_path = self._save_bytes_to_temp_file(content)
        else:
            raise ValueError(f"Unsupported content type: {type(content)}")

        source_info = source_info or {"source": "unknown", "source_type": "content"}

        # Update source info if we have a file path
        if isinstance(content, Path):
            source_info.update({
                "source": str(doc_path),
                "source_type": "file"
            })
        elif isinstance(content, str) and len(content) < 300 and '\n' not in content:
            try:
                if Path(content).exists():
                    source_info.update({
                        "source": str(doc_path),
                        "source_type": "file"
                    })
            except (OSError, ValueError):
                # Ignore file check errors for very long strings
                pass

        # Detect document format
        format_type = self._detect_document_format(doc_path)

        # Extract document content
        document_data = self._extract_document_content(doc_path, format_type)

        # Choose chunking strategy
        if self.chunk_by == "paragraphs":
            chunks = self._chunk_by_paragraphs(document_data, source_info)
        elif self.chunk_by == "sections":
            chunks = self._chunk_by_sections(document_data, source_info)
        elif self.chunk_by == "headings":
            chunks = self._chunk_by_headings(document_data, source_info)
        elif self.chunk_by == "pages":
            chunks = self._chunk_by_pages(document_data, source_info)
        elif self.chunk_by == "content_size":
            chunks = self._chunk_by_content_size(document_data, source_info)
        else:
            raise ValueError(f"Unknown chunking strategy: {self.chunk_by}")

        processing_time = time.time() - start_time

        return ChunkingResult(
            chunks=chunks,
            processing_time=processing_time,
            source_info={
                "document_format": format_type,
                "chunking_method": self.chunk_by,
                "backend_used": self.backend,
                "paragraphs_extracted": document_data.get("paragraph_count", 0),
                "tables_extracted": len(document_data.get("tables", [])),
                **source_info
            },
            strategy_used="doc_chunker"
        )

    def _detect_document_format(self, file_path: Path) -> str:
        """Detect the document format based on file extension."""
        suffix = file_path.suffix.lower()
        format_map = {
            '.docx': 'docx',
            '.doc': 'doc',
            '.odt': 'odt',
            '.rtf': 'rtf'
        }
        return format_map.get(suffix, 'unknown')

    def _save_content_to_temp_file(self, content: str) -> Path:
        """Save string content to temporary file."""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as tmp:
            tmp.write(content)
            return Path(tmp.name)

    def _save_bytes_to_temp_file(self, content: bytes) -> Path:
        """Save bytes content to temporary file."""
        with tempfile.NamedTemporaryFile(mode='wb', suffix='.doc', delete=False) as tmp:
            tmp.write(content)
            return Path(tmp.name)

    def _extract_document_content(self, file_path: Path, format_type: str) -> Dict[str, Any]:
        """Extract content from document using appropriate backend."""
        if format_type == 'docx' and self.backend == "python-docx" and HAS_PYTHON_DOCX:
            return self._extract_with_python_docx(file_path)
        elif format_type in ['docx', 'doc'] and self.backend == "mammoth" and HAS_MAMMOTH:
            return self._extract_with_mammoth(file_path)
        elif format_type == 'odt' and self.backend == "odf" and HAS_ODF:
            return self._extract_with_odf(file_path)
        elif format_type == 'rtf' and self.backend == "striprtf" and HAS_STRIPRTF:
            return self._extract_with_striprtf(file_path)
        elif self.backend == "textract" and HAS_TEXTRACT:
            return self._extract_with_textract(file_path)
        else:
            # Fallback to any available backend
            return self._extract_with_fallback(file_path, format_type)

    def _extract_with_python_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract content using python-docx library."""
        try:
            doc = DocxDocument(str(file_path))

            paragraphs = []
            tables = []
            headings = []

            # Extract paragraphs and identify headings
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    # Check if paragraph is a heading
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                        headings.append({
                            "text": para_text,
                            "level": level,
                            "paragraph_index": len(paragraphs)
                        })

                    paragraphs.append({
                        "text": para_text,
                        "style": para.style.name,
                        "is_heading": para.style.name.startswith('Heading'),
                        "formatting": self._extract_paragraph_formatting(para) if self.preserve_formatting else {}
                    })

            # Extract tables
            if self.extract_tables:
                for table_idx, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)

                    if table_data:
                        tables.append({
                            "index": table_idx,
                            "data": table_data,
                            "text": self._format_table_as_text(table_data)
                        })

            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "headings": headings,
                "paragraph_count": len(paragraphs),
                "backend": "python-docx"
            }

        except Exception as e:
            self.logger.error(f"Error extracting with python-docx: {e}")
            return self._extract_with_fallback(file_path, "docx")

    def _extract_paragraph_formatting(self, paragraph) -> Dict[str, Any]:
        """Extract formatting information from a paragraph."""
        formatting = {}

        if paragraph.runs:
            # Get formatting from first run as representative
            run = paragraph.runs[0]
            formatting.update({
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": run.font.size.pt if run.font.size else None
            })

        return formatting

    def _extract_with_mammoth(self, file_path: Path) -> Dict[str, Any]:
        """Extract content using mammoth library."""
        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                text = result.value

                # Split into paragraphs
                paragraphs = []
                for para_text in text.split('\n\n'):
                    para_text = para_text.strip()
                    if para_text:
                        paragraphs.append({
                            "text": para_text,
                            "style": "Normal",
                            "is_heading": False,
                            "formatting": {}
                        })

                return {
                    "paragraphs": paragraphs,
                    "tables": [],
                    "headings": [],
                    "paragraph_count": len(paragraphs),
                    "backend": "mammoth"
                }

        except Exception as e:
            self.logger.error(f"Error extracting with mammoth: {e}")
            return self._extract_with_fallback(file_path, "docx")

    def _extract_with_odf(self, file_path: Path) -> Dict[str, Any]:
        """Extract content using odfpy library."""
        try:
            doc = odf_load(str(file_path))
            paragraphs = []

            # Extract all paragraph elements
            for paragraph in doc.getElementsByType(ODFParagraph):
                para_text = teletype.extractText(paragraph).strip()
                if para_text:
                    paragraphs.append({
                        "text": para_text,
                        "style": "Normal",
                        "is_heading": False,
                        "formatting": {}
                    })

            return {
                "paragraphs": paragraphs,
                "tables": [],
                "headings": [],
                "paragraph_count": len(paragraphs),
                "backend": "odfpy"
            }

        except Exception as e:
            self.logger.error(f"Error extracting with odfpy: {e}")
            return self._extract_with_fallback(file_path, "odt")

    def _extract_with_striprtf(self, file_path: Path) -> Dict[str, Any]:
        """Extract content using striprtf library."""
        try:
            with open(file_path, 'r', encoding='utf-8') as rtf_file:
                rtf_content = rtf_file.read()
                text = rtf_to_text(rtf_content)

                # Split into paragraphs
                paragraphs = []
                for para_text in text.split('\n\n'):
                    para_text = para_text.strip()
                    if para_text:
                        paragraphs.append({
                            "text": para_text,
                            "style": "Normal",
                            "is_heading": False,
                            "formatting": {}
                        })

                return {
                    "paragraphs": paragraphs,
                    "tables": [],
                    "headings": [],
                    "paragraph_count": len(paragraphs),
                    "backend": "striprtf"
                }

        except Exception as e:
            self.logger.error(f"Error extracting with striprtf: {e}")
            return self._extract_with_fallback(file_path, "rtf")

    def _extract_with_textract(self, file_path: Path) -> Dict[str, Any]:
        """Extract content using textract library."""
        try:
            text = textract.process(str(file_path)).decode('utf-8')

            # Split into paragraphs
            paragraphs = []
            for para_text in text.split('\n\n'):
                para_text = para_text.strip()
                if para_text:
                    paragraphs.append({
                        "text": para_text,
                        "style": "Normal",
                        "is_heading": False,
                        "formatting": {}
                    })

            return {
                "paragraphs": paragraphs,
                "tables": [],
                "headings": [],
                "paragraph_count": len(paragraphs),
                "backend": "textract"
            }

        except Exception as e:
            self.logger.error(f"Error extracting with textract: {e}")
            return self._extract_with_fallback(file_path, "unknown")

    def _extract_with_fallback(self, file_path: Path, format_type: str) -> Dict[str, Any]:
        """Fallback content extraction method."""
        self.logger.warning(f"Using fallback extraction for {format_type} file")

        try:
            # Try to read as plain text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        except UnicodeDecodeError:
            # If that fails, try with different encoding
            with open(file_path, 'r', encoding='latin-1', errors='ignore') as f:
                text = f.read()

        # Basic paragraph splitting
        paragraphs = []
        for para_text in text.split('\n\n'):
            para_text = para_text.strip()
            if para_text and len(para_text) > 10:  # Filter out very short lines
                paragraphs.append({
                    "text": para_text,
                    "style": "Normal",
                    "is_heading": False,
                    "formatting": {}
                })

        return {
            "paragraphs": paragraphs,
            "tables": [],
            "headings": [],
            "paragraph_count": len(paragraphs),
            "backend": "fallback"
        }

    def _format_table_as_text(self, table_data: List[List[str]]) -> str:
        """Format table data as text."""
        if not table_data:
            return ""

        # Calculate column widths
        col_widths = []
        for col_idx in range(len(table_data[0])):
            max_width = max(len(str(row[col_idx])) if col_idx < len(row) else 0 for row in table_data)
            col_widths.append(max(max_width, 3))  # Minimum width of 3

        # Format table
        formatted_rows = []
        for row in table_data:
            formatted_cells = []
            for col_idx, cell in enumerate(row):
                if col_idx < len(col_widths):
                    formatted_cells.append(str(cell).ljust(col_widths[col_idx]))
            formatted_rows.append(" | ".join(formatted_cells))

        return "\n".join(formatted_rows)

    def _chunk_by_paragraphs(self, document_data: Dict[str, Any], source_info: Dict[str, Any]) -> List[Chunk]:
        """Chunk by grouping paragraphs."""
        chunks = []
        paragraphs = document_data.get("paragraphs", [])

        for i in range(0, len(paragraphs), self.paragraphs_per_chunk):
            chunk_paragraphs = paragraphs[i:i + self.paragraphs_per_chunk]
            content = "\n\n".join([p["text"] for p in chunk_paragraphs])

            if len(content.strip()) >= self.min_chunk_size:
                chunk = self._create_paragraph_chunk(chunk_paragraphs, content, i // self.paragraphs_per_chunk, source_info)
                chunks.append(chunk)

        # Add table chunks
        for table in document_data.get("tables", []):
            table_chunk = self._create_table_chunk(table, source_info)
            chunks.append(table_chunk)

        return chunks

    def _chunk_by_sections(self, document_data: Dict[str, Any], source_info: Dict[str, Any]) -> List[Chunk]:
        """Chunk by document sections (based on headings)."""
        chunks = []
        paragraphs = document_data.get("paragraphs", [])
        headings = document_data.get("headings", [])

        if not headings:
            # No headings found, fall back to paragraph chunking
            return self._chunk_by_paragraphs(document_data, source_info)

        # Create sections based on headings
        current_section = []
        current_heading = None
        section_index = 0

        for i, paragraph in enumerate(paragraphs):
            # Check if this paragraph is a heading
            paragraph_heading = next((h for h in headings if h["paragraph_index"] == i), None)

            if paragraph_heading and paragraph_heading["level"] in self.heading_levels:
                # Save previous section if it exists
                if current_section:
                    content = "\n\n".join([p["text"] for p in current_section])
                    if len(content.strip()) >= self.min_chunk_size:
                        chunk = self._create_section_chunk(current_section, content, current_heading, section_index, source_info)
                        chunks.append(chunk)
                        section_index += 1

                # Start new section
                current_section = [paragraph]
                current_heading = paragraph_heading
            else:
                current_section.append(paragraph)

        # Add final section
        if current_section:
            content = "\n\n".join([p["text"] for p in current_section])
            if len(content.strip()) >= self.min_chunk_size:
                chunk = self._create_section_chunk(current_section, content, current_heading, section_index, source_info)
                chunks.append(chunk)

        # Add table chunks
        for table in document_data.get("tables", []):
            table_chunk = self._create_table_chunk(table, source_info)
            chunks.append(table_chunk)

        return chunks

    def _chunk_by_headings(self, document_data: Dict[str, Any], source_info: Dict[str, Any]) -> List[Chunk]:
        """Chunk by heading hierarchy."""
        # This is similar to sections but respects all heading levels
        return self._chunk_by_sections(document_data, source_info)

    def _chunk_by_pages(self, document_data: Dict[str, Any], source_info: Dict[str, Any]) -> List[Chunk]:
        """Chunk by estimated pages (paragraph-based approximation)."""
        # Estimate paragraphs per page (rough approximation)
        estimated_paragraphs_per_page = 10
        return self._chunk_by_paragraphs(document_data, source_info)

    def _chunk_by_content_size(self, document_data: Dict[str, Any], source_info: Dict[str, Any]) -> List[Chunk]:
        """Chunk by content size while respecting paragraph boundaries."""
        chunks = []
        paragraphs = document_data.get("paragraphs", [])

        current_chunk = []
        current_size = 0
        chunk_index = 0

        for paragraph in paragraphs:
            para_text = paragraph["text"]
            para_size = len(para_text)

            # Check if adding this paragraph would exceed max size
            if current_size + para_size > self.max_chunk_size and current_chunk:
                # Create chunk from current content
                content = "\n\n".join([p["text"] for p in current_chunk])
                if len(content.strip()) >= self.min_chunk_size:
                    chunk = self._create_content_size_chunk(current_chunk, content, chunk_index, source_info)
                    chunks.append(chunk)
                    chunk_index += 1

                # Start new chunk
                current_chunk = [paragraph]
                current_size = para_size
            else:
                current_chunk.append(paragraph)
                current_size += para_size

        # Add final chunk
        if current_chunk:
            content = "\n\n".join([p["text"] for p in current_chunk])
            if len(content.strip()) >= self.min_chunk_size:
                chunk = self._create_content_size_chunk(current_chunk, content, chunk_index, source_info)
                chunks.append(chunk)

        # Add table chunks
        for table in document_data.get("tables", []):
            table_chunk = self._create_table_chunk(table, source_info)
            chunks.append(table_chunk)

        return chunks

    def _create_paragraph_chunk(self, paragraphs: List[Dict], content: str, chunk_index: int, source_info: Dict[str, Any]) -> Chunk:
        """Create a paragraph-based chunk."""
        metadata = ChunkMetadata(
            source=source_info.get("source", "unknown"),
            source_type=source_info.get("source_type", "content"),
            position=f"paragraph group {chunk_index + 1}",
            length=len(content),
            extra={
                "chunk_type": "paragraphs",
                "paragraph_count": len(paragraphs),
                "paragraph_styles": [p.get("style", "Normal") for p in paragraphs],
                "chunking_strategy": "paragraphs",
                "chunk_index": chunk_index
            }
        )

        return Chunk(
            id=f"doc_paragraphs_{chunk_index}",
            content=content,
            metadata=metadata,
            modality=ModalityType.TEXT
        )

    def _create_section_chunk(self, paragraphs: List[Dict], content: str, heading: Optional[Dict], section_index: int, source_info: Dict[str, Any]) -> Chunk:
        """Create a section-based chunk."""
        metadata = ChunkMetadata(
            source=source_info.get("source", "unknown"),
            source_type=source_info.get("source_type", "content"),
            position=f"section {section_index + 1}",
            length=len(content),
            extra={
                "chunk_type": "section",
                "section_heading": heading["text"] if heading else "Untitled Section",
                "heading_level": heading["level"] if heading else 0,
                "paragraph_count": len(paragraphs),
                "chunking_strategy": "sections",
                "section_index": section_index
            }
        )

        return Chunk(
            id=f"doc_section_{section_index}",
            content=content,
            metadata=metadata,
            modality=ModalityType.TEXT
        )

    def _create_content_size_chunk(self, paragraphs: List[Dict], content: str, chunk_index: int, source_info: Dict[str, Any]) -> Chunk:
        """Create a content size-based chunk."""
        metadata = ChunkMetadata(
            source=source_info.get("source", "unknown"),
            source_type=source_info.get("source_type", "content"),
            position=f"content chunk {chunk_index + 1}",
            length=len(content),
            extra={
                "chunk_type": "content_size",
                "paragraph_count": len(paragraphs),
                "chunking_strategy": "content_size",
                "chunk_index": chunk_index
            }
        )

        return Chunk(
            id=f"doc_content_{chunk_index}",
            content=content,
            metadata=metadata,
            modality=ModalityType.TEXT
        )

    def _create_table_chunk(self, table: Dict, source_info: Dict[str, Any]) -> Chunk:
        """Create a table chunk."""
        metadata = ChunkMetadata(
            source=source_info.get("source", "unknown"),
            source_type=source_info.get("source_type", "content"),
            position=f"table {table['index'] + 1}",
            length=len(table["text"]),
            extra={
                "chunk_type": "table",
                "table_index": table["index"],
                "row_count": len(table["data"]),
                "column_count": len(table["data"][0]) if table["data"] else 0,
                "chunking_strategy": "table"
            }
        )

        return Chunk(
            id=f"doc_table_{table['index']}",
            content=table["text"],
            metadata=metadata,
            modality=ModalityType.TABLE
        )

    # Streaming and adaptation methods
    def can_stream(self) -> bool:
        """Check if this chunker supports streaming."""
        return True

    def chunk_stream(self, content_stream, source_info=None, **kwargs):
        """Stream chunks as content becomes available."""
        # For documents, we need the full content to properly parse
        full_content = b""
        for chunk in content_stream:
            if isinstance(chunk, str):
                full_content += chunk.encode('utf-8')
            elif isinstance(chunk, bytes):
                full_content += chunk

        result = self.chunk(full_content, source_info=source_info, **kwargs)
        for chunk in result.chunks:
            yield chunk

    def adapt_parameters(self, feedback_score: float, feedback_type: str = "quality", **kwargs) -> None:
        """Adapt chunking parameters based on feedback."""
        adaptation_record = {
            "timestamp": time.time(),
            "feedback_score": feedback_score,
            "feedback_type": feedback_type,
            "old_config": {
                "paragraphs_per_chunk": self.paragraphs_per_chunk,
                "max_chunk_size": self.max_chunk_size
            }
        }

        # Apply adaptations based on feedback score
        if feedback_score < 0.5:  # Poor performance
            if feedback_type == "quality":
                # Reduce chunk size for better granularity
                self.paragraphs_per_chunk = max(1, int(self.paragraphs_per_chunk * 0.8))
                self.max_chunk_size = max(500, int(self.max_chunk_size * 0.8))
            elif feedback_type == "performance":
                # Increase chunk size for better performance
                self.paragraphs_per_chunk = min(20, int(self.paragraphs_per_chunk * 1.2))
                self.max_chunk_size = min(20000, int(self.max_chunk_size * 1.2))
        elif feedback_score > 0.8:  # Good performance
            if feedback_type == "quality":
                self.paragraphs_per_chunk = min(15, int(self.paragraphs_per_chunk * 1.1))
                self.max_chunk_size = min(15000, int(self.max_chunk_size * 1.1))

        # Handle specific feedback
        if kwargs.get("chunks_too_large"):
            self.paragraphs_per_chunk = max(1, int(self.paragraphs_per_chunk * 0.7))
            self.max_chunk_size = max(500, int(self.max_chunk_size * 0.7))
        elif kwargs.get("chunks_too_small"):
            self.paragraphs_per_chunk = min(20, int(self.paragraphs_per_chunk * 1.3))
            self.max_chunk_size = min(20000, int(self.max_chunk_size * 1.3))

        # Record new config
        adaptation_record["new_config"] = {
            "paragraphs_per_chunk": self.paragraphs_per_chunk,
            "max_chunk_size": self.max_chunk_size
        }

        # Store adaptation history
        if not hasattr(self, '_adaptation_history'):
            self._adaptation_history = []
        self._adaptation_history.append(adaptation_record)

        self.logger.info(f"Adapted parameters based on {feedback_type} feedback (score: {feedback_score})")

    def get_adaptation_history(self) -> List[Dict[str, Any]]:
        """Get history of adaptations made."""
        if not hasattr(self, '_adaptation_history'):
            self._adaptation_history = []
        return self._adaptation_history.copy()
